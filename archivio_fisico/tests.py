import datetime
import io
import shutil
import tempfile
import unittest

from django.contrib.auth import get_user_model
from django.core.files.base import ContentFile
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from django.utils import timezone

from anagrafiche.models import Anagrafica, Cliente
from archivio_fisico.forms import RigaOperazioneArchivioForm
from archivio_fisico.models import (
	CollocazioneFisica,
	OperazioneArchivio,
	RigaOperazioneArchivio,
	UnitaFisica,
	VerbaleConsegnaTemplate,
)
from archivio_fisico.word_templates import build_verbale_context, DocxTemplate as WordDocxTemplate
from documenti.models import Documento, DocumentiTipo
from fascicoli.models import Fascicolo
from protocollo.models import MovimentoProtocollo
from titolario.models import TitolarioVoce

try:  # pragma: no cover - optional dependency for docx generation tests
	from docx import Document as DocxDocument  # type: ignore
except ImportError:  # pragma: no cover
	DocxDocument = None


class UnitaListViewTests(TestCase):
	def setUp(self):
		self.user = get_user_model().objects.create_user("tester", password="password123")

	def _create_unita_chain(self):
		ufficio = UnitaFisica.objects.create(
			prefisso_codice="UFF",
			nome="Ufficio 1",
			tipo=UnitaFisica.Tipo.UFFICIO,
		)
		stanza = UnitaFisica.objects.create(
			prefisso_codice="ST",
			nome="Stanza 1",
			tipo=UnitaFisica.Tipo.STANZA,
			parent=ufficio,
		)
		scaffale = UnitaFisica.objects.create(
			prefisso_codice="SC",
			nome="Scaffale 1",
			tipo=UnitaFisica.Tipo.SCAFFALE,
			parent=stanza,
		)
		ripiano = UnitaFisica.objects.create(
			prefisso_codice="RP",
			nome="Ripiano 1",
			tipo=UnitaFisica.Tipo.RIPIANO,
			parent=scaffale,
		)
		return UnitaFisica.objects.create(
			prefisso_codice="CT",
			nome="Contenitore 1",
			tipo=UnitaFisica.Tipo.CONTENITORE,
			parent=ripiano,
		)

	def test_progressivo_calcolato(self):
		unita = self._create_unita_chain()
		atteso = f"{unita.codice}-{unita.nome}-{unita.get_tipo_display()}"
		self.assertEqual(unita.progressivo, atteso)

	def test_codice_progressivo_incrementale(self):
		ufficio = UnitaFisica.objects.create(prefisso_codice="UFF", nome="Ufficio", tipo=UnitaFisica.Tipo.UFFICIO)
		prima = UnitaFisica.objects.create(
			prefisso_codice="STZ",
			nome="Stanza 1",
			tipo=UnitaFisica.Tipo.STANZA,
			parent=ufficio,
		)
		seconda = UnitaFisica.objects.create(
			prefisso_codice="STZ",
			nome="Stanza 2",
			tipo=UnitaFisica.Tipo.STANZA,
			parent=ufficio,
		)
		self.assertEqual(prima.codice, "STZ1")
		self.assertEqual(seconda.codice, "STZ2")

	def _create_anagrafica_e_cliente(self):
		anagrafica = Anagrafica.objects.create(
			tipo=Anagrafica.TipoSoggetto.PERSONA_FISICA,
			nome="Mario",
			cognome="Rossi",
			codice_fiscale="RSSMRA80A01H501U",
		)
		cliente = Cliente.objects.create(anagrafica=anagrafica)
		return anagrafica, cliente

	def test_clienti_tree_present_in_context(self):
		self.client.force_login(self.user)
		response = self.client.get(reverse("archivio_fisico:unita_list"))
		self.assertEqual(response.status_code, 200)
		self.assertIn("clienti_tree", response.context)
		self.assertEqual(response.context["clienti_tree"], [])

	def test_clienti_tree_with_corrente_data(self):
		self.client.force_login(self.user)
		temp_dir = tempfile.mkdtemp()
		self.addCleanup(lambda: shutil.rmtree(temp_dir, ignore_errors=True))

		with override_settings(ARCHIVIO_BASE_PATH=temp_dir):
			contenitore = self._create_unita_chain()
			anagrafica, cliente = self._create_anagrafica_e_cliente()

			voce = TitolarioVoce.objects.create(codice="01", titolo="Amministrazione")
			documento_tipo = DocumentiTipo.objects.create(codice="TEST", nome="Documento test")

			fascicolo = Fascicolo.objects.create(
				cliente=cliente,
				titolario_voce=voce,
				anno=datetime.date.today().year,
				titolo="Fascicolo corrente",
				ubicazione=contenitore,
			)

			today = datetime.date.today()
			documento_in_fascicolo = Documento.objects.create(
				tipo=documento_tipo,
				fascicolo=fascicolo,
				cliente=cliente,
				titolario_voce=voce,
				descrizione="Documento collegato",
				data_documento=today,
				stato=Documento.Stato.DEFINITIVO,
			)
			documento_sciolto = Documento.objects.create(
				tipo=documento_tipo,
				cliente=cliente,
				titolario_voce=voce,
				descrizione="Documento sciolto",
				data_documento=today,
				stato=Documento.Stato.DEFINITIVO,
			)

			CollocazioneFisica.objects.create(documento=documento_in_fascicolo, unita=contenitore, attiva=True)
			CollocazioneFisica.objects.create(documento=documento_sciolto, unita=contenitore, attiva=True)

			response = self.client.get(reverse("archivio_fisico:unita_list"))

		self.assertEqual(response.status_code, 200)
		clienti_tree = response.context["clienti_tree"]
		self.assertEqual(len(clienti_tree), 1)
		entry = clienti_tree[0]
		self.assertEqual(entry["label"], anagrafica.display_name())
		self.assertEqual(len(entry["fascicoli"]), 1)
		self.assertEqual(entry["documenti_count"], 2)
		self.assertEqual(len(entry["documenti_sciolti"]), 1)
		self.assertEqual(len(entry["fascicoli"][0].documenti_corrente), 1)

	def test_fascicolo_without_ubicazione_included_via_document_collocazione(self):
		self.client.force_login(self.user)
		temp_dir = tempfile.mkdtemp()
		self.addCleanup(lambda: shutil.rmtree(temp_dir, ignore_errors=True))

		with override_settings(ARCHIVIO_BASE_PATH=temp_dir):
			contenitore = self._create_unita_chain()
			anagrafica, cliente = self._create_anagrafica_e_cliente()
			voce = TitolarioVoce.objects.create(codice="02", titolo="Contabilità")
			documento_tipo = DocumentiTipo.objects.create(codice="LIB", nome="Libro giornale")

			fascicolo = Fascicolo.objects.create(
				cliente=cliente,
				titolario_voce=voce,
				anno=datetime.date.today().year,
				titolo="Fascicolo senza ubicazione",
			)

			documento = Documento.objects.create(
				tipo=documento_tipo,
				fascicolo=fascicolo,
				cliente=cliente,
				titolario_voce=voce,
				descrizione="Documento collegato",
				data_documento=datetime.date.today(),
				stato=Documento.Stato.DEFINITIVO,
			)

			CollocazioneFisica.objects.create(documento=documento, unita=contenitore, attiva=True)

			response = self.client.get(reverse("archivio_fisico:unita_list"))

		self.assertEqual(response.status_code, 200)
		clienti_tree = response.context["clienti_tree"]
		self.assertEqual(len(clienti_tree), 1)
		entry = clienti_tree[0]
		self.assertEqual(len(entry["fascicoli"]), 1)
		self.assertEqual(entry["documenti_count"], 1)
		self.assertEqual(len(entry["fascicoli"][0].documenti_corrente), 1)


class UnitaInventoryTreeViewTests(TestCase):
	def setUp(self):
		self.user = get_user_model().objects.create_user("inventory", password="password123")
		self.client.force_login(self.user)

	def _create_unit_chain(self):
		ufficio = UnitaFisica.objects.create(
			prefisso_codice="UF",
			nome="Ufficio Inventario",
			tipo=UnitaFisica.Tipo.UFFICIO,
		)
		stanza = UnitaFisica.objects.create(
			prefisso_codice="STI",
			nome="Stanza Inventario",
			tipo=UnitaFisica.Tipo.STANZA,
			parent=ufficio,
		)
		scaffale = UnitaFisica.objects.create(
			prefisso_codice="SCF",
			nome="Scaffale Inventario",
			tipo=UnitaFisica.Tipo.SCAFFALE,
			parent=stanza,
		)
		ripiano = UnitaFisica.objects.create(
			prefisso_codice="RPI",
			nome="Ripiano Inventario",
			tipo=UnitaFisica.Tipo.RIPIANO,
			parent=scaffale,
		)
		contenitore = UnitaFisica.objects.create(
			prefisso_codice="CTI",
			nome="Contenitore Inventario",
			tipo=UnitaFisica.Tipo.CONTENITORE,
			parent=ripiano,
		)
		return ufficio, contenitore

	def _create_anagrafica_cliente(self):
		anagrafica = Anagrafica.objects.create(
			tipo=Anagrafica.TipoSoggetto.PERSONA_FISICA,
			nome="Anna",
			cognome="Verdi",
			codice_fiscale="VRDANN80A01H501R",
		)
		cliente = Cliente.objects.create(anagrafica=anagrafica)
		return anagrafica, cliente

	def _find_node(self, nodes, unit_id):
		for node in nodes:
			if node["unit"].pk == unit_id:
				return node
			found = self._find_node(node.get("children", []), unit_id)
			if found:
				return found
		return None

	def test_inventory_tree_empty(self):
		response = self.client.get(reverse("archivio_fisico:unita_inventory_tree"))
		self.assertEqual(response.status_code, 200)
		self.assertIn("unita_tree", response.context)
		self.assertEqual(response.context["unita_tree"], [])
		self.assertIn("generated_at", response.context)

	def test_inventory_tree_with_fascicoli_and_documents(self):
		temp_dir = tempfile.mkdtemp()
		self.addCleanup(lambda: shutil.rmtree(temp_dir, ignore_errors=True))

		with override_settings(ARCHIVIO_BASE_PATH=temp_dir):
			root, contenitore = self._create_unit_chain()
			anagrafica, cliente = self._create_anagrafica_cliente()
			voce = TitolarioVoce.objects.create(codice="03", titolo="Inventario")
			documento_tipo = DocumentiTipo.objects.create(codice="INV", nome="Documento inventario")

			fascicolo = Fascicolo.objects.create(
				cliente=cliente,
				titolario_voce=voce,
				anno=datetime.date.today().year,
				titolo="Fascicolo inventario",
				ubicazione=contenitore,
			)

			CollocazioneFisica.objects.create(
				content_object=fascicolo,
				unita=contenitore,
				attiva=True,
			)

			documento_fascicolo = Documento.objects.create(
				tipo=documento_tipo,
				fascicolo=fascicolo,
				cliente=cliente,
				titolario_voce=voce,
				descrizione="Documento nel fascicolo",
				data_documento=datetime.date.today(),
				stato=Documento.Stato.DEFINITIVO,
				digitale=False,
			)
			CollocazioneFisica.objects.create(
				documento=documento_fascicolo,
				unita=contenitore,
				attiva=True,
			)

			documento_sciolto = Documento.objects.create(
				tipo=documento_tipo,
				cliente=cliente,
				titolario_voce=voce,
				descrizione="Documento sciolto inventario",
				data_documento=datetime.date.today(),
				stato=Documento.Stato.DEFINITIVO,
				digitale=False,
			)
			CollocazioneFisica.objects.create(
				documento=documento_sciolto,
				unita=contenitore,
				attiva=True,
			)

			response = self.client.get(reverse("archivio_fisico:unita_inventory_tree"))

			self.assertEqual(response.status_code, 200)
			unita_tree = response.context["unita_tree"]
			self.assertTrue(unita_tree)
			root_node = self._find_node(unita_tree, root.pk)
			self.assertIsNotNone(root_node)
			contenitore_node = self._find_node(unita_tree, contenitore.pk)
			self.assertIsNotNone(contenitore_node)
			self.assertEqual(len(contenitore_node["fascicoli"]), 1)
			fasc_entry = contenitore_node["fascicoli"][0]
			self.assertEqual(fasc_entry["fascicolo"].pk, fascicolo.pk)
			self.assertEqual(len(fasc_entry["documenti"]), 1)
			self.assertEqual(fasc_entry["documenti"][0].pk, documento_fascicolo.pk)
			self.assertEqual(len(contenitore_node["documenti_sciolti"]), 1)
			self.assertEqual(contenitore_node["documenti_sciolti"][0].pk, documento_sciolto.pk)


@unittest.skipIf(WordDocxTemplate is None or DocxDocument is None, "docxtpl non disponibile")
class VerbaleConsegnaDocxTests(TestCase):
	def setUp(self):
		self.user = get_user_model().objects.create_user("verbale", password="password123")
		self.client.force_login(self.user)

		temp_dir = tempfile.mkdtemp()
		self.addCleanup(lambda: shutil.rmtree(temp_dir, ignore_errors=True))
		settings_override = override_settings(ARCHIVIO_BASE_PATH=temp_dir, MEDIA_ROOT=temp_dir)
		settings_override.enable()
		self.addCleanup(settings_override.disable)

		self.anagrafica = Anagrafica.objects.create(
			tipo=Anagrafica.TipoSoggetto.PERSONA_FISICA,
			nome="Luigi",
			cognome="Bianchi",
			codice_fiscale="BNCLGU80A01F205Z",
		)
		self.cliente = Cliente.objects.create(anagrafica=self.anagrafica)
		self.voce = TitolarioVoce.objects.create(codice="01", titolo="Affari generali")
		self.documento_tipo = DocumentiTipo.objects.create(codice="VERB", nome="Verbale")

		self.unita_root = UnitaFisica.objects.create(
			prefisso_codice="UFF",
			nome="Sede centrale",
			tipo=UnitaFisica.Tipo.UFFICIO,
		)
		self.unita_dest = UnitaFisica.objects.create(
			prefisso_codice="ST",
			nome="Stanza archivio",
			tipo=UnitaFisica.Tipo.STANZA,
			parent=self.unita_root,
		)

		self.fascicolo = Fascicolo.objects.create(
			cliente=self.cliente,
			titolario_voce=self.voce,
			anno=timezone.now().year,
			titolo="Fascicolo consegna",
			ubicazione=self.unita_root,
		)

		self.documento = Documento.objects.create(
			tipo=self.documento_tipo,
			fascicolo=self.fascicolo,
			cliente=self.cliente,
			titolario_voce=self.voce,
			descrizione="Documento per verbale",
			data_documento=timezone.now().date(),
			stato=Documento.Stato.DEFINITIVO,
			digitale=False,
		)

		self.movimento = MovimentoProtocollo.objects.create(
			documento=self.documento,
			cliente=self.anagrafica,
			direzione="OUT",
			data=timezone.now(),
			anno=timezone.now().year,
			numero=1,
			operatore=self.user,
			destinatario="Destinatario test",
			causale="Consegna documenti",
			note="Protocollo di uscita",
		)

		self.operazione = OperazioneArchivio.objects.create(
			tipo_operazione="uscita",
			referente_interno=self.user,
			referente_esterno=self.anagrafica,
			note="Operazione di consegna",
		)

		RigaOperazioneArchivio.objects.create(
			operazione=self.operazione,
			documento=self.documento,
			stato_precedente="In archivio",
			stato_successivo="Consegnato",
			note="Riga di test",
			movimento_protocollo=self.movimento,
			unita_fisica_sorgente=self.unita_root,
			unita_fisica_destinazione=self.unita_dest,
		)

		self.template = self._create_template()

	def _create_template(self) -> VerbaleConsegnaTemplate:
		doc = DocxDocument()
		doc.add_paragraph("Verbale {{ operazione.tipo_label }}")
		doc.add_paragraph("Documento {{ righe[0].documento.codice }}")
		buffer = io.BytesIO()
		doc.save(buffer)
		buffer.seek(0)
		template = VerbaleConsegnaTemplate(
			nome="Verbale base",
			is_default=True,
		)
		template.file_template.save("verbale_base.docx", ContentFile(buffer.getvalue()), save=False)
		template.save()
		return template

	def test_build_verbale_context_includes_protocollo(self):
		context = build_verbale_context(self.operazione)
		self.assertEqual(context["operazione"]["tipo_label"], "Uscita")
		self.assertEqual(len(context["righe"]), 1)
		riga = context["righe"][0]
		self.assertEqual(riga["documento"]["codice"], self.documento.codice)
		self.assertEqual(
			riga["protocollo"]["protocollo"],
			f"{self.movimento.anno}/{self.movimento.numero:06d}",
		)
		self.assertEqual(riga["unita_sorgente"]["id"], self.unita_root.pk)
		self.assertEqual(riga["unita_destinazione"]["id"], self.unita_dest.pk)
		self.assertEqual(context["unita"]["sorgente"]["id"], self.unita_root.pk)
		self.assertEqual(context["unita"]["destinazione"]["id"], self.unita_dest.pk)

	def test_download_view_returns_word_document(self):
		url = reverse("archivio_fisico:operazionearchivio_verbale_docx", args=[self.operazione.pk])
		response = self.client.get(url)
		self.assertEqual(response.status_code, 200)
		self.assertEqual(
			response["Content-Type"],
			"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
		)
		self.assertIn(".docx", response["Content-Disposition"])

		rendered = DocxDocument(io.BytesIO(response.content))
		full_text = "\n".join(p.text for p in rendered.paragraphs)
		self.assertIn(self.operazione.get_tipo_operazione_display(), full_text)
		self.assertIn(self.documento.codice, full_text)


class OperazioneArchivioViewTests(TestCase):
	def setUp(self):
		self.user = get_user_model().objects.create_user("operatore", password="password123")
		self.client.force_login(self.user)

		temp_dir = tempfile.mkdtemp()
		self.addCleanup(lambda: shutil.rmtree(temp_dir, ignore_errors=True))
		settings_override = override_settings(ARCHIVIO_BASE_PATH=temp_dir, MEDIA_ROOT=temp_dir)
		settings_override.enable()
		self.addCleanup(settings_override.disable)

		self.anagrafica = Anagrafica.objects.create(
			tipo=Anagrafica.TipoSoggetto.PERSONA_FISICA,
			nome="Mario",
			cognome="Rossi",
			codice_fiscale="RSSMRA80A01H501U",
		)
		self.titolario = TitolarioVoce.objects.create(codice="10", titolo="Archivio")
		self.documento_tipo = DocumentiTipo.objects.create(codice="DOC", nome="Documento generico")
		self.unita_sorgente = UnitaFisica.objects.create(
			prefisso_codice="UFF",
			nome="Ufficio Test",
			tipo=UnitaFisica.Tipo.UFFICIO,
		)
		self.unita_destinazione = UnitaFisica.objects.create(
			prefisso_codice="ST",
			nome="Stanza Test",
			tipo=UnitaFisica.Tipo.STANZA,
			parent=self.unita_sorgente,
		)
		self.cliente_record = Cliente.objects.create(anagrafica=self.anagrafica)

		scarico_override = override_settings(ARCHIVIO_FISICO_UNITA_SCARICO_ID=self.unita_destinazione.pk)
		scarico_override.enable()
		self.addCleanup(scarico_override.disable)

		self.documento = Documento.objects.create(
			tipo=self.documento_tipo,
			cliente=self.cliente_record,
			titolario_voce=self.titolario,
			descrizione="Documento movimentato",
			data_documento=datetime.date.today(),
			stato=Documento.Stato.DEFINITIVO,
			digitale=False,
		)
		CollocazioneFisica.objects.create(
			content_object=self.documento,
			documento=self.documento,
			unita=self.unita_sorgente,
			dal=timezone.now().date(),
		)
		self.documento_secondario = Documento.objects.create(
			tipo=self.documento_tipo,
			cliente=self.cliente_record,
			titolario_voce=self.titolario,
			descrizione="Documento collegato",
			data_documento=datetime.date.today(),
			stato=Documento.Stato.DEFINITIVO,
			digitale=False,
		)
		CollocazioneFisica.objects.create(
			content_object=self.documento_secondario,
			documento=self.documento_secondario,
			unita=self.unita_sorgente,
			dal=timezone.now().date(),
		)
		self.documento_senza_storico = Documento.objects.create(
			tipo=self.documento_tipo,
			cliente=self.cliente_record,
			titolario_voce=self.titolario,
			descrizione="Documento senza operazioni archivio",
			data_documento=datetime.date.today(),
			stato=Documento.Stato.DEFINITIVO,
			digitale=False,
		)
		CollocazioneFisica.objects.create(
			content_object=self.documento_senza_storico,
			documento=self.documento_senza_storico,
			unita=self.unita_sorgente,
			dal=timezone.now().date(),
		)
		self.movimento = MovimentoProtocollo.registra_entrata(
			documento=self.documento,
			operatore=self.user,
			ubicazione=self.unita_sorgente,
			da_chi="Ufficio test",
		)
		self.movimento_secondario = MovimentoProtocollo.registra_entrata(
			documento=self.documento_secondario,
			operatore=self.user,
			ubicazione=self.unita_sorgente,
			da_chi="Ufficio test",
		)
		self.movimento_documento_senza_storico = MovimentoProtocollo.registra_entrata(
			documento=self.documento_senza_storico,
			operatore=self.user,
			ubicazione=self.unita_sorgente,
			da_chi="Ufficio test",
		)
		operazione_entrata_documento = OperazioneArchivio.objects.create(
			tipo_operazione="entrata",
			referente_interno=self.user,
			referente_esterno=self.anagrafica,
		)
		RigaOperazioneArchivio.objects.create(
			operazione=operazione_entrata_documento,
			documento=self.documento,
		)
		operazione_entrata_documento_secondario = OperazioneArchivio.objects.create(
			tipo_operazione="entrata",
			referente_interno=self.user,
			referente_esterno=self.anagrafica,
		)
		RigaOperazioneArchivio.objects.create(
			operazione=operazione_entrata_documento_secondario,
			documento=self.documento_secondario,
		)
		self.documento_fuori = Documento.objects.create(
			tipo=self.documento_tipo,
			cliente=self.cliente_record,
			titolario_voce=self.titolario,
			descrizione="Documento fuori",
			data_documento=datetime.date.today(),
			stato=Documento.Stato.DEFINITIVO,
			digitale=False,
		)
		CollocazioneFisica.objects.create(
			content_object=self.documento_fuori,
			documento=self.documento_fuori,
			unita=self.unita_sorgente,
			dal=timezone.now().date(),
		)
		self.movimento_documento_fuori = MovimentoProtocollo.objects.create(
			documento=self.documento_fuori,
			cliente=self.anagrafica,
			direzione="OUT",
			data=timezone.now(),
			anno=timezone.now().year,
			numero=5000,
			operatore=self.user,
			ubicazione=self.unita_sorgente,
			chiuso=False,
		)
		operazione_entrata_documento_fuori = OperazioneArchivio.objects.create(
			tipo_operazione="entrata",
			referente_interno=self.user,
			referente_esterno=self.anagrafica,
		)
		RigaOperazioneArchivio.objects.create(
			operazione=operazione_entrata_documento_fuori,
			documento=self.documento_fuori,
		)
		operazione_uscita_documento_fuori = OperazioneArchivio.objects.create(
			tipo_operazione="uscita",
			referente_interno=self.user,
			referente_esterno=self.anagrafica,
		)
		RigaOperazioneArchivio.objects.create(
			operazione=operazione_uscita_documento_fuori,
			documento=self.documento_fuori,
		)

		current_year = timezone.now().year
		self.fascicolo_in = Fascicolo.objects.create(
			cliente=self.cliente_record,
			titolario_voce=self.titolario,
			anno=current_year,
			titolo="Fascicolo in archivio",
			ubicazione=self.unita_sorgente,
		)
		self.movimento_fascicolo_in = MovimentoProtocollo.objects.create(
			fascicolo=self.fascicolo_in,
			cliente=self.anagrafica,
			direzione="IN",
			data=timezone.now(),
			anno=current_year,
			numero=5001,
			operatore=self.user,
			ubicazione=self.unita_sorgente,
			chiuso=True,
		)
		self.fascicolo_fuori = Fascicolo.objects.create(
			cliente=self.cliente_record,
			titolario_voce=self.titolario,
			anno=current_year,
			titolo="Fascicolo fuori",
			ubicazione=self.unita_sorgente,
		)
		self.movimento_fascicolo_out = MovimentoProtocollo.objects.create(
			fascicolo=self.fascicolo_fuori,
			cliente=self.anagrafica,
			direzione="OUT",
			data=timezone.now(),
			anno=current_year,
			numero=5002,
			operatore=self.user,
			ubicazione=self.unita_sorgente,
			chiuso=False,
		)
		operazione_entrata_fascicolo = OperazioneArchivio.objects.create(
			tipo_operazione="entrata",
			referente_interno=self.user,
			referente_esterno=self.anagrafica,
		)
		RigaOperazioneArchivio.objects.create(
			operazione=operazione_entrata_fascicolo,
			fascicolo=self.fascicolo_in,
		)
		operazione_entrata_fascicolo_fuori = OperazioneArchivio.objects.create(
			tipo_operazione="entrata",
			referente_interno=self.user,
			referente_esterno=self.anagrafica,
		)
		RigaOperazioneArchivio.objects.create(
			operazione=operazione_entrata_fascicolo_fuori,
			fascicolo=self.fascicolo_fuori,
		)
		operazione_uscita_fascicolo_fuori = OperazioneArchivio.objects.create(
			tipo_operazione="uscita",
			referente_interno=self.user,
			referente_esterno=self.anagrafica,
		)
		RigaOperazioneArchivio.objects.create(
			operazione=operazione_uscita_fascicolo_fuori,
			fascicolo=self.fascicolo_fuori,
		)

	def _build_create_payload(self):
		return {
			"tipo_operazione": "uscita",
			"referente_interno": str(self.user.pk),
			"referente_esterno": str(self.anagrafica.pk),
			"note": "Operazione di test",
			"righe-TOTAL_FORMS": "1",
			"righe-INITIAL_FORMS": "0",
			"righe-MIN_NUM_FORMS": "0",
			"righe-MAX_NUM_FORMS": "1000",
			"righe-0-id": "",
			"righe-0-fascicolo": "",
			"righe-0-documento": str(self.documento.pk),
			"righe-0-stato_precedente": Documento.Stato.DEFINITIVO,
			"righe-0-stato_successivo": Documento.Stato.SCARICATO,
			"righe-0-note": "Riga test",
			"righe-0-movimento_protocollo": str(self.movimento.pk),
			"righe-0-unita_fisica_sorgente": str(self.unita_sorgente.pk),
			"righe-0-unita_fisica_destinazione": str(self.unita_destinazione.pk),
		}

	def test_operazione_create_flow(self):
		url = reverse("archivio_fisico:operazionearchivio_create")
		response = self.client.post(url, data=self._build_create_payload())
		self.assertEqual(response.status_code, 302)
		operazione = OperazioneArchivio.objects.latest("id")
		self.assertEqual(operazione.referente_interno, self.user)
		righe = list(operazione.righe.all())
		self.assertEqual(len(righe), 1)
		self.assertEqual(righe[0].documento, self.documento)
		self.assertEqual(righe[0].stato_successivo, Documento.Stato.SCARICATO)
		self.assertEqual(righe[0].movimento_protocollo, self.movimento)
		self.assertEqual(righe[0].unita_fisica_sorgente, self.unita_sorgente)
		self.assertEqual(righe[0].unita_fisica_destinazione, self.unita_destinazione)
		self.documento.refresh_from_db()
		self.movimento.refresh_from_db()
		self.assertEqual(self.documento.stato, Documento.Stato.SCARICATO)
		self.assertEqual(self.movimento.ubicazione, self.unita_sorgente)

	def test_operazione_create_with_verbale_scan(self):
		url = reverse("archivio_fisico:operazionearchivio_create")
		payload = self._build_create_payload()
		payload["verbale_scan"] = SimpleUploadedFile(
			"verbale.pdf",
			b"%PDF-1.4 test file",
			content_type="application/pdf",
		)
		response = self.client.post(url, data=payload)
		self.assertEqual(response.status_code, 302)
		operazione = OperazioneArchivio.objects.latest("id")
		self.assertTrue(operazione.verbale_scan)
		self.assertTrue(operazione.verbale_scan.name.startswith("archivio/operazioni/"))
		self.assertTrue(operazione.verbale_scan.name.endswith(".pdf"))

	def test_operazione_uscita_documento_senza_storico_archivio(self):
		url = reverse("archivio_fisico:operazionearchivio_create")
		payload = self._build_create_payload()
		payload.update(
			{
				"righe-0-documento": str(self.documento_senza_storico.pk),
				"righe-0-movimento_protocollo": str(self.movimento_documento_senza_storico.pk),
			}
		)
		response = self.client.post(url, data=payload)
		self.assertEqual(response.status_code, 302)
		operazione = OperazioneArchivio.objects.latest("id")
		riga = operazione.righe.get()
		self.assertEqual(riga.documento, self.documento_senza_storico)
		self.assertEqual(riga.unita_fisica_sorgente, self.unita_sorgente)
		self.assertEqual(riga.unita_fisica_destinazione, self.unita_destinazione)
		self.documento_senza_storico.refresh_from_db()
		self.assertEqual(self.documento_senza_storico.stato, Documento.Stato.SCARICATO)

	def test_operazione_update_flow(self):
		operazione = OperazioneArchivio.objects.create(
			tipo_operazione="interna",
			referente_interno=self.user,
			referente_esterno=self.anagrafica,
			note="Nota iniziale",
		)
		riga = RigaOperazioneArchivio.objects.create(
			operazione=operazione,
			documento=self.documento,
			stato_precedente="In archivio",
			stato_successivo=Documento.Stato.DEFINITIVO,
			note="Riga iniziale",
			movimento_protocollo=self.movimento,
			unita_fisica_sorgente=self.unita_sorgente,
			unita_fisica_destinazione=self.unita_destinazione,
		)

		url = reverse("archivio_fisico:operazionearchivio_update", args=[operazione.pk])
		payload = {
			"tipo_operazione": "interna",
			"referente_interno": str(self.user.pk),
			"referente_esterno": str(self.anagrafica.pk),
			"note": "Nota aggiornata",
			"righe-TOTAL_FORMS": "1",
			"righe-INITIAL_FORMS": "1",
			"righe-MIN_NUM_FORMS": "0",
			"righe-MAX_NUM_FORMS": "1000",
			"righe-0-id": str(riga.pk),
			"righe-0-fascicolo": "",
			"righe-0-documento": str(self.documento.pk),
			"righe-0-stato_precedente": Documento.Stato.DEFINITIVO,
			"righe-0-stato_successivo": Documento.Stato.CONSEGNATO,
			"righe-0-note": "Riga aggiornata",
			"righe-0-DELETE": "",
			"righe-0-movimento_protocollo": str(self.movimento.pk),
			"righe-0-unita_fisica_sorgente": str(self.unita_sorgente.pk),
			"righe-0-unita_fisica_destinazione": str(self.unita_destinazione.pk),
		}
		response = self.client.post(url, data=payload)
		self.assertEqual(response.status_code, 302)
		operazione.refresh_from_db()
		self.assertEqual(operazione.note, "Nota aggiornata")
		riga.refresh_from_db()
		self.assertEqual(riga.stato_successivo, Documento.Stato.CONSEGNATO)
		self.assertEqual(riga.note, "Riga aggiornata")
		self.assertEqual(riga.movimento_protocollo, self.movimento)
		self.assertEqual(riga.unita_fisica_sorgente, self.unita_sorgente)
		self.assertEqual(riga.unita_fisica_destinazione, self.unita_destinazione)
		self.documento.refresh_from_db()
		self.movimento.refresh_from_db()
		self.assertEqual(self.documento.stato, Documento.Stato.CONSEGNATO)
		self.assertEqual(self.movimento.ubicazione, self.unita_destinazione)

	def test_operazione_create_with_multiple_righe(self):
		url = reverse("archivio_fisico:operazionearchivio_create")
		payload = self._build_create_payload()
		payload.update(
			{
				"righe-TOTAL_FORMS": "2",
				"righe-1-id": "",
				"righe-1-fascicolo": "",
				"righe-1-documento": str(self.documento_secondario.pk),
				"righe-1-stato_precedente": Documento.Stato.DEFINITIVO,
				"righe-1-stato_successivo": Documento.Stato.SCARICATO,
				"righe-1-note": "Riga aggiuntiva",
				"righe-1-movimento_protocollo": str(self.movimento_secondario.pk),
				"righe-1-unita_fisica_sorgente": str(self.unita_sorgente.pk),
				"righe-1-unita_fisica_destinazione": str(self.unita_destinazione.pk),
			}
		)
		response = self.client.post(url, data=payload)
		self.assertEqual(response.status_code, 302)
		operazione = OperazioneArchivio.objects.latest("id")
		righe = list(operazione.righe.order_by("id"))
		self.assertEqual(len(righe), 2)
		self.assertEqual(
			{r.documento_id for r in righe},
			{self.documento.pk, self.documento_secondario.pk},
		)
		self.assertEqual(righe[0].movimento_protocollo, self.movimento)
		self.assertEqual(righe[1].movimento_protocollo, self.movimento_secondario)
		for riga in righe:
			self.assertEqual(riga.unita_fisica_sorgente, self.unita_sorgente)
			self.assertEqual(riga.unita_fisica_destinazione, self.unita_destinazione)
		self.documento.refresh_from_db()
		self.documento_secondario.refresh_from_db()
		self.movimento.refresh_from_db()
		self.movimento_secondario.refresh_from_db()
		self.assertEqual(self.documento.stato, Documento.Stato.SCARICATO)
		self.assertEqual(self.documento_secondario.stato, Documento.Stato.SCARICATO)
		self.assertEqual(self.movimento.ubicazione, self.unita_sorgente)
		self.assertEqual(self.movimento_secondario.ubicazione, self.unita_sorgente)

	def test_riga_form_filters_documenti_per_giacenza(self):
		form_entrata = RigaOperazioneArchivioForm(
			instance=RigaOperazioneArchivio(),
			tipo_operazione="entrata",
			referente_esterno_id=self.anagrafica.pk,
		)
		entrata_ids = set(form_entrata.fields["documento"].queryset.values_list("pk", flat=True))
		self.assertIn(self.documento_fuori.pk, entrata_ids)
		self.assertNotIn(self.documento.pk, entrata_ids)
		self.assertNotIn(self.documento_secondario.pk, entrata_ids)

		form_uscita = RigaOperazioneArchivioForm(
			instance=RigaOperazioneArchivio(),
			tipo_operazione="uscita",
			referente_esterno_id=self.anagrafica.pk,
		)
		uscita_ids = set(form_uscita.fields["documento"].queryset.values_list("pk", flat=True))
		self.assertIn(self.documento.pk, uscita_ids)
		self.assertIn(self.documento_secondario.pk, uscita_ids)
		self.assertIn(self.documento_senza_storico.pk, uscita_ids)
		self.assertNotIn(self.documento_fuori.pk, uscita_ids)

		form_interna = RigaOperazioneArchivioForm(
			instance=RigaOperazioneArchivio(),
			tipo_operazione="interna",
			referente_esterno_id=self.anagrafica.pk,
		)
		interna_ids = set(form_interna.fields["documento"].queryset.values_list("pk", flat=True))
		self.assertEqual(uscita_ids, interna_ids)
		self.assertIn(self.documento_senza_storico.pk, interna_ids)
		self.assertNotIn(self.documento_fuori.pk, interna_ids)

	def test_riga_form_filters_fascicoli_per_giacenza(self):
		form_entrata = RigaOperazioneArchivioForm(
			instance=RigaOperazioneArchivio(),
			tipo_operazione="entrata",
			referente_esterno_id=self.anagrafica.pk,
		)
		entrata_ids = set(form_entrata.fields["fascicolo"].queryset.values_list("pk", flat=True))
		self.assertIn(self.fascicolo_fuori.pk, entrata_ids)
		self.assertNotIn(self.fascicolo_in.pk, entrata_ids)

		form_uscita = RigaOperazioneArchivioForm(
			instance=RigaOperazioneArchivio(),
			tipo_operazione="uscita",
			referente_esterno_id=self.anagrafica.pk,
		)
		uscita_ids = set(form_uscita.fields["fascicolo"].queryset.values_list("pk", flat=True))
		self.assertIn(self.fascicolo_in.pk, uscita_ids)
		self.assertNotIn(self.fascicolo_fuori.pk, uscita_ids)

		form_interna = RigaOperazioneArchivioForm(
			instance=RigaOperazioneArchivio(),
			tipo_operazione="interna",
			referente_esterno_id=self.anagrafica.pk,
		)
		interna_ids = set(form_interna.fields["fascicolo"].queryset.values_list("pk", flat=True))
		self.assertEqual(uscita_ids, interna_ids)
		self.assertNotIn(self.fascicolo_fuori.pk, interna_ids)
