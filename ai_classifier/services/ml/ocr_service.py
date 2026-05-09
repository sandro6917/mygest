"""
OCR Service - Estrazione testo da PDF e immagini

Gestisce:
- PDF con testo nativo (pdfplumber)
- PDF scansionati (OCR con pytesseract)
- Immagini (OCR con pytesseract)
- File ZIP (estrae e processa contenuti)
- Fallback automatico se testo nativo è vuoto
"""
import logging
import os
import tempfile
import shutil
import zipfile
from typing import Optional, Dict, Any, List
from pathlib import Path

# PDF Processing
import pdfplumber
from pdf2image import convert_from_path
from PyPDF2 import PdfReader

# OCR
import pytesseract
from PIL import Image

logger = logging.getLogger(__name__)


class OCRService:
    """
    Servizio per estrazione testo da documenti con supporto OCR.
    """
    
    def __init__(
        self,
        tesseract_lang: str = 'ita+eng',
        ocr_dpi: int = 300,
        max_pages_ocr: int = 10,
        min_text_length: int = 100,
    ):
        """
        Inizializza OCR Service.
        
        Args:
            tesseract_lang: Lingue Tesseract (default: italiano + inglese)
            ocr_dpi: DPI per conversione PDF → immagine
            max_pages_ocr: Numero massimo pagine da processare con OCR
            min_text_length: Lunghezza minima testo per considerarlo valido
        """
        self.tesseract_lang = tesseract_lang
        self.ocr_dpi = ocr_dpi
        self.max_pages_ocr = max_pages_ocr
        self.min_text_length = min_text_length
        
        # Verifica che tesseract sia installato
        try:
            pytesseract.get_tesseract_version()
            logger.info(f"✅ Tesseract OCR disponibile (lang: {tesseract_lang})")
        except Exception as e:
            logger.warning(f"⚠️ Tesseract OCR non trovato: {e}")
    
    def extract_text_from_file(
        self,
        file_path: str,
        mime_type: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Estrae testo da file con gestione automatica del tipo.
        
        Args:
            file_path: Path al file
            mime_type: MIME type (opzionale, verrà auto-rilevato)
            
        Returns:
            Dict con:
            {
                'text': str,
                'method': 'native' | 'ocr' | 'hybrid',
                'pages': int,
                'confidence': float (se OCR),
                'metadata': dict
            }
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File non trovato: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        # ZIP - Estrai e processa contenuti
        if file_ext == '.zip':
            return self._extract_from_zip(file_path)
        
        # PDF
        if file_ext == '.pdf' or (mime_type and 'pdf' in mime_type):
            return self._extract_from_pdf(file_path)
        
        # Immagini
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp']:
            return self._extract_from_image(file_path)
        
        # Office (DOCX, XLSX)
        elif file_ext == '.docx':
            return self._extract_from_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self._extract_from_excel(file_path)
        
        # Text files
        elif file_ext in ['.txt', '.text', '.log', '.csv']:
            return self._extract_from_text(file_path)
        
        else:
            logger.warning(f"Tipo file non supportato per OCR: {file_ext}")
            return {
                'text': '',
                'method': 'unsupported',
                'pages': 0,
                'metadata': {'error': f'Tipo file non supportato: {file_ext}'}
            }
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Estrae testo da PDF.
        Prova prima con testo nativo, poi OCR se necessario.
        """
        logger.info(f"📄 Elaborazione PDF: {file_path}")
        
        # Tentativo 1: Testo nativo con pdfplumber
        text_native = ""
        num_pages = 0
        
        try:
            with pdfplumber.open(file_path) as pdf:
                num_pages = len(pdf.pages)
                logger.info(f"  📖 PDF ha {num_pages} pagine")
                
                # Estrai testo da tutte le pagine (max 50 per performance)
                max_pages = min(num_pages, 50)
                for i, page in enumerate(pdf.pages[:max_pages]):
                    page_text = page.extract_text() or ''
                    text_native += page_text + "\n"
                
                logger.info(f"  ✅ Estratto testo nativo: {len(text_native)} caratteri")
        
        except Exception as e:
            logger.warning(f"  ⚠️ Errore estrazione testo nativo: {e}")
        
        # Se testo nativo sufficiente, ritorna
        if len(text_native.strip()) >= self.min_text_length:
            # Estrai metadata PDF
            metadata = self._extract_pdf_metadata(file_path)
            
            return {
                'text': text_native.strip(),
                'method': 'native',
                'pages': num_pages,
                'metadata': metadata
            }
        
        # Tentativo 2: OCR (PDF scansionato)
        logger.info(f"  🔍 Testo nativo insufficiente ({len(text_native)} char), provo OCR...")
        
        try:
            text_ocr = self._ocr_pdf_to_text(file_path)
            metadata = self._extract_pdf_metadata(file_path)
            
            # Combina testo nativo + OCR (hybrid)
            combined_text = text_native + "\n" + text_ocr
            
            return {
                'text': combined_text.strip(),
                'method': 'hybrid' if text_native else 'ocr',
                'pages': num_pages,
                'metadata': metadata
            }
        
        except Exception as e:
            logger.error(f"  ❌ Errore OCR PDF: {e}")
            return {
                'text': text_native.strip(),
                'method': 'native_fallback',
                'pages': num_pages,
                'metadata': {'ocr_error': str(e)}
            }
    
    def _ocr_pdf_to_text(self, file_path: str) -> str:
        """
        Converte PDF in immagini e applica OCR.
        """
        text_parts = []
        
        # Converti PDF → immagini
        try:
            images = convert_from_path(
                file_path,
                dpi=self.ocr_dpi,
                first_page=1,
                last_page=self.max_pages_ocr,
            )
            
            logger.info(f"  🖼️  Convertite {len(images)} pagine in immagini")
            
            # OCR su ogni immagine
            for i, image in enumerate(images, 1):
                logger.info(f"  🔍 OCR pagina {i}/{len(images)}...")
                page_text = pytesseract.image_to_string(
                    image,
                    lang=self.tesseract_lang,
                )
                text_parts.append(page_text)
            
            full_text = "\n".join(text_parts)
            logger.info(f"  ✅ OCR completato: {len(full_text)} caratteri")
            return full_text
        
        except Exception as e:
            logger.error(f"  ❌ Errore conversione PDF → immagini: {e}")
            raise
    
    def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """
        Estrae testo da immagine con OCR.
        """
        logger.info(f"🖼️  Elaborazione immagine: {file_path}")
        
        try:
            image = Image.open(file_path)
            
            # OCR
            text = pytesseract.image_to_string(
                image,
                lang=self.tesseract_lang,
            )
            
            # Ottieni confidence (se disponibile)
            ocr_data = pytesseract.image_to_data(
                image,
                lang=self.tesseract_lang,
                output_type=pytesseract.Output.DICT
            )
            
            # Calcola confidence media
            confidences = [
                float(conf) for conf in ocr_data['conf']
                if conf != '-1'
            ]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            logger.info(f"  ✅ OCR immagine: {len(text)} caratteri (confidence: {avg_confidence:.1f}%)")
            
            return {
                'text': text.strip(),
                'method': 'ocr',
                'pages': 1,
                'confidence': avg_confidence / 100.0,  # 0-1
                'metadata': {
                    'image_size': image.size,
                    'image_format': image.format,
                }
            }
        
        except Exception as e:
            logger.error(f"  ❌ Errore OCR immagine: {e}")
            return {
                'text': '',
                'method': 'ocr_failed',
                'pages': 1,
                'metadata': {'error': str(e)}
            }
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Estrae testo da file DOCX.
        """
        logger.info(f"📝 Elaborazione DOCX: {file_path}")
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Estrai testo da tutti i paragrafi
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            full_text = "\n".join(text_parts)
            
            # Estrai anche testo da tabelle
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text += "\n" + row_text
            
            logger.info(f"  ✅ Testo estratto da DOCX: {len(full_text)} caratteri")
            
            return {
                'text': full_text.strip(),
                'method': 'native',
                'pages': 1,  # DOCX non ha concept di pagine
                'metadata': {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                }
            }
        
        except Exception as e:
            logger.error(f"  ❌ Errore estrazione DOCX: {e}")
            return {
                'text': '',
                'method': 'native_failed',
                'pages': 1,
                'metadata': {'error': str(e)}
            }
    
    def _extract_from_excel(self, file_path: str) -> Dict[str, Any]:
        """
        Estrae testo da file Excel (XLSX/XLS).
        """
        logger.info(f"📊 Elaborazione Excel: {file_path}")
        
        try:
            import openpyxl
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_parts = []
            
            # Estrai testo da tutti i fogli
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Header foglio
                text_parts.append(f"\n=== {sheet_name} ===\n")
                
                # Celle con valore
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(
                        str(cell) for cell in row
                        if cell is not None and str(cell).strip()
                    )
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            logger.info(f"  ✅ Testo estratto da Excel: {len(full_text)} caratteri")
            
            return {
                'text': full_text.strip(),
                'method': 'native',
                'pages': len(workbook.sheetnames),
                'metadata': {
                    'sheets': workbook.sheetnames,
                }
            }
        
        except Exception as e:
            logger.error(f"  ❌ Errore estrazione Excel: {e}")
            return {
                'text': '',
                'method': 'native_failed',
                'pages': 1,
                'metadata': {'error': str(e)}
            }
    
    def _extract_from_zip(self, file_path: str) -> Dict[str, Any]:
        """
        Estrae e processa tutti i file contenuti in un archivio ZIP.
        
        Returns:
            Dict con testo aggregato di tutti i file processabili
        """
        logger.info(f"📦 Elaborazione ZIP: {file_path}")
        
        temp_dir = None
        all_texts = []
        processed_files = []
        total_pages = 0
        methods_used = set()
        
        try:
            # Crea directory temporanea
            temp_dir = tempfile.mkdtemp(prefix='mygest_zip_')
            
            # Estrai ZIP
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                logger.info(f"  📂 Trovati {len(file_list)} file nello ZIP")
                
                # Estrai tutti i file
                zip_ref.extractall(temp_dir)
            
            # Processa ogni file estratto
            for file_name in file_list:
                # Salta directory e file nascosti
                if file_name.endswith('/') or file_name.startswith('.') or '/__MACOSX' in file_name:
                    continue
                
                extracted_path = os.path.join(temp_dir, file_name)
                
                # Verifica che il file esista ed è un file regolare
                if not os.path.isfile(extracted_path):
                    continue
                
                logger.info(f"  📄 Processando: {file_name}")
                
                try:
                    # Chiama ricorsivamente extract_text_from_file
                    result = self.extract_text_from_file(extracted_path)
                    
                    if result and result.get('text'):
                        text = result['text'].strip()
                        
                        # Salta file con testo troppo corto
                        if len(text) >= 50:  # Soglia minima per file in ZIP
                            all_texts.append(f"=== {file_name} ===\n{text}")
                            processed_files.append(file_name)
                            total_pages += result.get('pages', 1)
                            methods_used.add(result.get('method', 'unknown'))
                            logger.info(f"    ✅ Estratti {len(text)} caratteri")
                        else:
                            logger.info(f"    ⚠️ Testo troppo corto ({len(text)} char), skip")
                    
                except Exception as e:
                    logger.warning(f"    ⚠️ Errore processando {file_name}: {e}")
                    continue
            
            # Combina tutti i testi estratti
            combined_text = "\n\n".join(all_texts)
            
            logger.info(f"  ✅ ZIP processato: {len(processed_files)}/{len(file_list)} file estratti, "
                       f"{len(combined_text)} caratteri totali")
            
            return {
                'text': combined_text,
                'method': 'zip_aggregate',
                'pages': total_pages,
                'metadata': {
                    'zip_file': os.path.basename(file_path),
                    'total_files': len(file_list),
                    'processed_files': processed_files,
                    'methods_used': list(methods_used),
                }
            }
        
        except zipfile.BadZipFile:
            logger.error(f"  ❌ File ZIP corrotto o non valido: {file_path}")
            return {
                'text': '',
                'method': 'zip_failed',
                'pages': 0,
                'metadata': {'error': 'File ZIP corrotto'}
            }
        
        except Exception as e:
            logger.error(f"  ❌ Errore processando ZIP: {e}")
            return {
                'text': '',
                'method': 'zip_failed',
                'pages': 0,
                'metadata': {'error': str(e)}
            }
        
        finally:
            # Pulizia directory temporanea
            if temp_dir and os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                    logger.info(f"  🧹 Directory temporanea pulita")
                except Exception as e:
                    logger.warning(f"  ⚠️ Errore pulizia temp dir: {e}")
    
    def _extract_from_text(self, file_path: str) -> Dict[str, Any]:
        """
        Estrae testo da file di testo plain (.txt, .csv, .log, etc.).
        """
        logger.info(f"📝 Elaborazione file testo: {file_path}")
        
        try:
            # Prova varie codifiche comuni
            encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
            text = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    encoding_used = encoding
                    break
                except (UnicodeDecodeError, LookupError):
                    continue
            
            if text is None:
                logger.error(f"  ❌ Impossibile leggere il file con nessuna codifica")
                return {
                    'text': '',
                    'method': 'text_failed',
                    'pages': 1,
                    'metadata': {'error': 'Encoding non supportato'}
                }
            
            logger.info(f"  ✅ Testo estratto: {len(text)} caratteri (encoding: {encoding_used})")
            
            return {
                'text': text.strip(),
                'method': 'text',
                'pages': 1,
                'metadata': {
                    'encoding': encoding_used,
                    'lines': text.count('\n') + 1,
                }
            }
        
        except Exception as e:
            logger.error(f"  ❌ Errore lettura file testo: {e}")
            return {
                'text': '',
                'method': 'text_failed',
                'pages': 1,
                'metadata': {'error': str(e)}
            }
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Estrae metadata da PDF (autore, data creazione, etc.).
        """
        metadata = {}
        
        try:
            reader = PdfReader(file_path)
            pdf_meta = reader.metadata
            
            if pdf_meta:
                metadata = {
                    'author': pdf_meta.get('/Author', ''),
                    'creator': pdf_meta.get('/Creator', ''),
                    'producer': pdf_meta.get('/Producer', ''),
                    'subject': pdf_meta.get('/Subject', ''),
                    'title': pdf_meta.get('/Title', ''),
                    'creation_date': str(pdf_meta.get('/CreationDate', '')),
                    'mod_date': str(pdf_meta.get('/ModDate', '')),
                }
        
        except Exception as e:
            logger.warning(f"Errore estrazione metadata PDF: {e}")
        
        return metadata
